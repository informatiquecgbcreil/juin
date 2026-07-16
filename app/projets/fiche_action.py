"""Fiche action au format imposé : schéma, saisie et génération du document.

Une « fiche action » est le document structuré (format imposé par la
fédération) décrivant une action du projet social. Son contenu narratif
est saisi une seule fois et stocké en JSON sur ``ProjetAction.fiche_json``.
Le générateur produit le document Word officiel, en y injectant
automatiquement les indicateurs quantitatifs calculés par l'application
(séances, présences, participants uniques…) : plus de recopie.

Deux formats de section :
- ``liste``  : une puce par ligne saisie ;
- ``table``  : une ligne par ligne saisie, colonnes séparées par « | ».
"""
from __future__ import annotations

import json
from dataclasses import dataclass

# En-tête (bloc d'identité en haut de la fiche).
CHAMPS_ENTETE = [
    ("axe", "Axe"),
    ("public_principal", "Public principal"),
    ("public_associe", "Public associé"),
    ("echelle", "Échelle"),
]


@dataclass(frozen=True)
class Section:
    cle: str
    numero: int
    titre: str
    format: str            # "liste" | "table"
    colonnes: tuple = ()   # pour les tables


# Trame imposée (ordre et intitulés calqués sur le document fédéral).
SECTIONS: list[Section] = [
    Section("objectif_general", 1, "Intitulé et objectif général", "liste"),
    Section("objectifs_specifiques", 2, "Objectifs spécifiques concernés", "table",
            ("Objectif", "Contribution de l'action")),
    Section("porte_entree", 3, "Porte d'entrée habitant", "liste"),
    Section("besoins", 4, "Besoins repérés", "liste"),
    Section("public_vise", 5, "Public visé", "table", ("Public principal", "Public associé")),
    Section("deroule", 6, "Déroulé concret de l'action", "table", ("Étape", "Description concrète")),
    Section("posture", 7, "Posture professionnelle mobilisée", "liste"),
    Section("parcours", 8, "Parcours habitant dans l'action", "table",
            ("Étape", "Ce que vit l'habitant", "Ce que fait le secteur")),
    Section("role_poste", 9, "Rôle du poste créé", "liste"),
    Section("transversalite", 10, "Transversalité interne", "table",
            ("Secteur / public", "Apport possible de l'action")),
    Section("portee", 11, "Portée au-delà du Centre", "liste"),
    Section("moyens", 12, "Moyens mobilisés", "table", ("Moyens humains", "Moyens matériels")),
    Section("resultats", 13, "Résultats attendus", "table", ("Temporalité", "Résultats attendus")),
    Section("indicateurs_quali", 14, "Indicateurs d'évaluation (qualitatifs)", "liste"),
    Section("vigilance", 15, "Points de vigilance", "liste"),
    Section("odd", 16, "Objectifs de développement durable concernés", "table",
            ("ODD", "Rattachement à l'action")),
]

SECTIONS_PAR_CLE = {s.cle: s for s in SECTIONS}


def charger(fiche_json: str | None) -> dict:
    """Désérialise le contenu de la fiche (tolérant aux valeurs vides/corrompues)."""
    if not fiche_json:
        return {}
    try:
        data = json.loads(fiche_json)
        return data if isinstance(data, dict) else {}
    except (ValueError, TypeError):
        return {}


def depuis_formulaire(form, existant: dict | None = None) -> dict:
    """Construit le dict de fiche depuis les champs du formulaire d'édition.

    Les champs NON rendus par l'éditeur (ex. la section 2 quand des objectifs
    sectoriels structurés la remplacent) sont ABSENTS du formulaire : on
    conserve alors leur valeur existante au lieu de l'écraser à vide.
    """
    existant = existant or {}
    ex_entete = existant.get("entete", {}) or {}
    ex_sections = existant.get("sections", {}) or {}

    data: dict = {"entete": {}, "sections": {}}
    for cle, _ in CHAMPS_ENTETE:
        champ = f"entete_{cle}"
        if champ in form:
            data["entete"][cle] = (form.get(champ) or "").strip()
        else:
            data["entete"][cle] = (ex_entete.get(cle) or "").strip()
    for s in SECTIONS:
        champ = f"section_{s.cle}"
        if champ in form:
            data["sections"][s.cle] = (form.get(champ) or "").strip()
        else:
            data["sections"][s.cle] = (ex_sections.get(s.cle) or "").strip()
    return data


def vers_json(data: dict) -> str:
    return json.dumps(data, ensure_ascii=False)


def _lignes(brut: str) -> list[str]:
    return [l.strip() for l in (brut or "").splitlines() if l.strip()]


def _table_rows(brut: str, n_cols: int) -> list[list[str]]:
    rows = []
    for ligne in _lignes(brut):
        cells = [c.strip() for c in ligne.split("|")]
        cells += [""] * (n_cols - len(cells))
        rows.append(cells[:n_cols])
    return rows


def indicateurs_quantitatifs(stats: dict) -> list[tuple[str, str]]:
    """Transforme les stats calculées de l'action en lignes (libellé, valeur)."""
    kpi = (stats or {}).get("kpi", {}) or {}
    demo = (stats or {}).get("demo", {}) or {}
    qpv = demo.get("qpv", {}) or {}
    return [
        ("Nombre de séances réalisées", str(kpi.get("sessions", 0) or 0)),
        ("Participations cumulées (présences)", str(kpi.get("presences", 0) or 0)),
        ("Participants uniques", str(kpi.get("uniques", 0) or 0)),
        ("Heures-participants", str(kpi.get("hours_people", 0) or 0)),
        ("Dont quartier prioritaire (QPV)", str(qpv.get("qpv", 0) or 0)),
        ("Âge moyen", str(demo.get("age_avg") if demo.get("age_avg") is not None else "—")),
    ]


def generer_docx(projet, action, fiche: dict, stats: dict, annee: int, objectifs: list | None = None):
    """Construit le document Word de la fiche action au format imposé.

    ``objectifs`` : contributions structurées [{objectif, contribution}] ;
    si fournies, elles alimentent la section 2 (sinon repli sur le texte libre).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    doc.add_heading(f"Fiche action — {action.titre}", level=0)
    sous = doc.add_paragraph(f"{projet.nom} · {projet.secteur} · année {annee}")
    sous.runs[0].italic = True

    entete = fiche.get("entete", {}) if fiche else {}
    sections = fiche.get("sections", {}) if fiche else {}

    # --- Bloc d'identité (tableau 2 colonnes) ---
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for cle, label in CHAMPS_ENTETE:
        valeur = (entete.get(cle) or "").strip()
        if valeur:
            row = t.add_row().cells
            row[0].text = label
            row[1].text = valeur
            row[0].paragraphs[0].runs[0].bold = True
    if not t.rows:
        # tableau vide : on le retire proprement
        t._element.getparent().remove(t._element)

    # --- Sections numérotées ---
    for s in SECTIONS:
        doc.add_heading(f"{s.numero}. {s.titre}", level=1)
        brut = sections.get(s.cle, "")

        # Section 2 : priorité aux objectifs sectoriels structurés (référentiel).
        if s.cle == "objectifs_specifiques" and objectifs:
            tab = doc.add_table(rows=1, cols=2)
            tab.style = "Light Grid Accent 1"
            for j, col in enumerate(("Objectif", "Contribution de l'action")):
                cell = tab.rows[0].cells[j]
                cell.text = col
                cell.paragraphs[0].runs[0].bold = True
            for item in objectifs:
                obj = item["objectif"]
                cells = tab.add_row().cells
                cells[0].text = f"{obj.code} — {obj.libelle}"
                cells[1].text = item.get("contribution") or ""
            continue

        if s.format == "liste":
            items = _lignes(brut)
            if items:
                for it in items:
                    doc.add_paragraph(it, style="List Bullet")
            else:
                _vide(doc)
        else:  # table
            rows = _table_rows(brut, len(s.colonnes))
            if rows:
                tab = doc.add_table(rows=1, cols=len(s.colonnes))
                tab.style = "Light Grid Accent 1"
                for j, col in enumerate(s.colonnes):
                    cell = tab.rows[0].cells[j]
                    cell.text = col
                    cell.paragraphs[0].runs[0].bold = True
                for r in rows:
                    cells = tab.add_row().cells
                    for j, val in enumerate(r):
                        cells[j].text = val
            else:
                _vide(doc)

        # Section 14 : on injecte les indicateurs quantitatifs RÉELS de l'app.
        if s.cle == "indicateurs_quali":
            p = doc.add_paragraph()
            run = p.add_run("Indicateurs quantitatifs (calculés automatiquement par l'application)")
            run.bold = True
            tab = doc.add_table(rows=0, cols=2)
            tab.style = "Light Grid Accent 1"
            for label, valeur in indicateurs_quantitatifs(stats):
                cells = tab.add_row().cells
                cells[0].text = label
                cells[1].text = valeur

    pied = doc.add_paragraph(
        "Document généré automatiquement — les indicateurs quantitatifs "
        "proviennent des données réelles de l'application."
    )
    pied.runs[0].italic = True
    pied.runs[0].font.size = Pt(8)
    return doc


def _vide(doc):
    p = doc.add_paragraph("(à compléter)")
    p.runs[0].italic = True
