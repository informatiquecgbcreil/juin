"""Location payante de salles et de matériel.

Une réservation de salle ou un prêt de matériel peut être une mise à
disposition facturée (``est_location``). Ce service gère le suivi du
paiement, la restitution de la caution, les listes de travail (à encaisser,
cautions à rendre), les recettes encaissées, et la génération d'une
convention / reçu de location au format Word.

Volontairement AUTONOME : rien n'est écrit dans le module Caisse ou la
comptabilité — on note seulement le règlement (montant, mode, date) sur la
location elle-même.
"""
from __future__ import annotations

from datetime import date

from app.extensions import db
from app.models import InventaireItem, PretMateriel, ReservationSalle


# ---------------------------------------------------------------------------
# Transitions de paiement / caution
# ---------------------------------------------------------------------------

def marquer_paye(obj, mode: str | None = None, jour: date | None = None) -> None:
    """Enregistre le règlement du loyer d'une location."""
    if not getattr(obj, "est_location", False):
        return
    obj.paiement_statut = "regle"
    if mode:
        obj.paiement_mode = (mode or "").strip() or None
    obj.paiement_date = jour or date.today()
    db.session.commit()


def annuler_paiement(obj) -> None:
    """Repasse une location « à régler » (correction de saisie)."""
    if not getattr(obj, "est_location", False):
        return
    obj.paiement_statut = "a_regler"
    obj.paiement_date = None
    db.session.commit()


def marquer_caution_rendue(obj) -> None:
    obj.caution_rendue = True
    db.session.commit()


# ---------------------------------------------------------------------------
# Listes de travail
# ---------------------------------------------------------------------------

def _reservations_location(secteur: str | None):
    q = ReservationSalle.query.filter(
        ReservationSalle.est_location.is_(True),
        ReservationSalle.statut == "approuvee",
    )
    if secteur:
        q = q.filter(db.or_(ReservationSalle.secteur == secteur, ReservationSalle.secteur.is_(None)))
    return q


def _prets_location(secteur: str | None):
    q = PretMateriel.query.filter(
        PretMateriel.est_location.is_(True),
        PretMateriel.statut == "approuvee",
    )
    if secteur:
        q = q.join(InventaireItem).filter(InventaireItem.secteur == secteur)
    return q


def locations_a_encaisser(secteur: str | None = None) -> dict:
    """Locations approuvées dont le loyer reste à régler."""
    reservations = (
        _reservations_location(secteur)
        .filter(ReservationSalle.paiement_statut == "a_regler")
        .order_by(ReservationSalle.date_reservation.asc())
        .all()
    )
    prets = (
        _prets_location(secteur)
        .filter(PretMateriel.paiement_statut == "a_regler")
        .order_by(PretMateriel.date_pret.asc())
        .all()
    )
    total = sum(float(r.tarif_montant or 0) for r in reservations) + \
        sum(float(p.tarif_montant or 0) for p in prets)
    return {"reservations": reservations, "prets": prets, "total": round(total, 2)}


def cautions_a_rendre(secteur: str | None = None) -> list:
    """Cautions encore détenues (location réglée ou non, caution non rendue)."""
    res = [r for r in _reservations_location(secteur).all() if r.caution_a_rendre]
    prets = [p for p in _prets_location(secteur).all() if p.caution_a_rendre]
    return res + prets


def recettes_encaissees(date_debut: date | None = None, date_fin: date | None = None,
                        secteur: str | None = None) -> dict:
    """Somme des loyers réglés (par date de paiement) sur la période."""
    def _dans_periode(obj) -> bool:
        d = obj.paiement_date
        if d is None:
            return False
        if date_debut and d < date_debut:
            return False
        if date_fin and d > date_fin:
            return False
        return True

    reservations = [
        r for r in _reservations_location(secteur).filter(ReservationSalle.paiement_statut == "regle").all()
        if _dans_periode(r)
    ]
    prets = [
        p for p in _prets_location(secteur).filter(PretMateriel.paiement_statut == "regle").all()
        if _dans_periode(p)
    ]
    total = sum(float(r.tarif_montant or 0) for r in reservations) + \
        sum(float(p.tarif_montant or 0) for p in prets)
    return {
        "reservations": reservations, "prets": prets,
        "nb": len(reservations) + len(prets), "total": round(total, 2),
    }


# ---------------------------------------------------------------------------
# Convention / reçu de location (DOCX)
# ---------------------------------------------------------------------------

def _org_nom() -> str:
    from flask import current_app

    from app.models import InstanceSettings
    row = InstanceSettings.query.first()
    if row:
        nom = (row.organization_name or "").strip() or (row.app_name or "").strip()
        if nom:
            return nom
    return current_app.config.get("APP_NAME") or "La structure"


def _fmt_euros(valeur) -> str:
    try:
        return f"{float(valeur or 0):.2f} €".replace(".", ",")
    except (TypeError, ValueError):
        return "0,00 €"


def _details_objet(obj) -> dict:
    """Normalise une réservation OU un prêt en éléments de convention."""
    if isinstance(obj, ReservationSalle):
        ressource = obj.salle.nom if obj.salle else "Salle"
        periode = (
            f"le {obj.date_reservation.strftime('%d/%m/%Y')} "
            f"de {obj.heure_debut} à {obj.heure_fin}"
        )
        return {
            "type": "Salle", "ressource": ressource, "periode": periode,
            "locataire": obj.locataire or "", "contact": obj.contact or "",
        }
    # PretMateriel
    ressource = obj.item.designation if obj.item else "Matériel"
    if obj.quantite and obj.quantite > 1:
        ressource += f" (×{obj.quantite})"
    periode = f"du {obj.date_pret.strftime('%d/%m/%Y')}"
    if obj.date_retour_prevue:
        periode += f" au {obj.date_retour_prevue.strftime('%d/%m/%Y')}"
    return {
        "type": "Matériel", "ressource": ressource, "periode": periode,
        "locataire": obj.emprunteur or "", "contact": obj.contact or "",
    }


def construire_convention_docx(obj):
    """Convention de mise à disposition / location, à relire et signer."""
    from docx import Document

    d = _details_objet(obj)
    org = _org_nom()
    doc = Document()

    doc.add_heading("Convention de location / mise à disposition", level=1)
    doc.add_paragraph(f"Entre {org} (le loueur)")
    doc.add_paragraph(
        f"et {d['locataire'] or '……………………………'} (le locataire)"
        + (f" — {d['contact']}" if d["contact"] else "")
    )

    doc.add_heading("Objet de la location", level=2)
    doc.add_paragraph(f"Nature : {d['type']} — {d['ressource']}")
    doc.add_paragraph(f"Période : {d['periode']}")
    if getattr(obj, "motif", None):
        doc.add_paragraph(f"Usage prévu : {obj.motif}")

    doc.add_heading("Conditions financières", level=2)
    montant = _fmt_euros(obj.tarif_montant)
    base = f" ({obj.tarif_unite_label})" if getattr(obj, "tarif_unite", None) else ""
    doc.add_paragraph(f"Montant de la location : {montant}{base}")
    if obj.caution_montant:
        doc.add_paragraph(
            f"Dépôt de garantie (caution) : {_fmt_euros(obj.caution_montant)} — "
            "restitué après état des lieux / retour du matériel sans dommage."
        )
    statut = "réglé" if obj.est_paye else "à régler"
    ligne_regl = f"Règlement : {statut}"
    if obj.est_paye and obj.paiement_date:
        ligne_regl += f" le {obj.paiement_date.strftime('%d/%m/%Y')}"
    doc.add_paragraph(ligne_regl)

    doc.add_heading("Engagements du locataire", level=2)
    for point in (
        "Utiliser les lieux / le matériel conformément à leur destination.",
        "Respecter les règles de sécurité et le règlement intérieur.",
        "Signaler et prendre en charge toute dégradation constatée.",
        "Restituer les lieux / le matériel dans leur état d'origine.",
    ):
        doc.add_paragraph(point, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        f"Fait le {date.today().strftime('%d/%m/%Y')}. "
        "Signatures (le loueur / le locataire) :"
    )
    doc.add_paragraph("\n\n\n____________________            ____________________")
    doc.add_paragraph(
        "Document généré automatiquement par l'application — à relire, "
        "compléter et faire signer avant la mise à disposition."
    )
    return doc
