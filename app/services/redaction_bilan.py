"""Aide à la rédaction du bilan financeur — trame pré-remplie.

Toutes les données d'un bilan financeur existent déjà dans l'application :
montants et lignes de la subvention, ateliers financés (et donc séances,
heures, participations, publics via l'émargement), projets liés (journal,
indicateurs), narratif annuel du secteur. Ce service les assemble en une
TRAME : des paragraphes pré-rédigés, chiffrés et sourcés, que l'équipe
copie puis ajuste — jamais de texte inventé, uniquement des formulations
types remplies avec les chiffres réels.
"""
from __future__ import annotations

from datetime import date

from app.extensions import db
from app.models import (
    BilanLourdNarratif,
    ProjetJournalEntry,
    Subvention,
)


def _fmt_euros(valeur) -> str:
    entier = f"{float(valeur or 0):,.0f}".replace(",", " ")
    return f"{entier} €"


def _fmt_nombre(valeur) -> str:
    if valeur is None:
        return "—"
    f = float(valeur)
    if f == int(f):
        return f"{int(f):,}".replace(",", " ")
    return f"{f:,.1f}".replace(",", " ").replace(".", ",")


def _stats_activite(subvention: Subvention, annee: int) -> dict:
    """Activité réelle des ateliers financés sur l'exercice (émargement)."""
    from app.services.transitions import _stats_ateliers

    atelier_ids = [
        lien.atelier_id for lien in (subvention.ateliers_finances or [])
        if lien.atelier is not None and not lien.atelier.is_deleted
    ]
    stats = _stats_ateliers(atelier_ids, date(annee, 1, 1), date(annee, 12, 31))
    stats["atelier_ids"] = atelier_ids
    return stats


def _demographie(atelier_ids: list[int], annee: int) -> dict | None:
    """Démographie des publics touchés (nécessite un contexte utilisateur)."""
    if not atelier_ids:
        return None
    try:
        from app.statsimpact.engine import StatsFilters, compute_demography_stats

        flt = StatsFilters(
            atelier_ids=atelier_ids,
            date_from=date(annee, 1, 1),
            date_to=date(annee, 12, 31),
        )
        return compute_demography_stats(flt)
    except Exception:  # hors requête (tests unitaires purs) : section omise
        return None


def _finances(subvention: Subvention) -> dict:
    lignes = []
    for ligne in subvention.lignes or []:
        if getattr(ligne, "nature", "charge") != "charge":
            continue
        lignes.append({
            "compte": ligne.compte,
            "libelle": ligne.libelle,
            "budget": round(float(ligne.montant_reel or 0), 2),
            "engage": round(float(ligne.engage or 0), 2),
            "reste": round(float(ligne.reste or 0), 2),
        })
    return {
        "lignes": lignes,
        "total_budget": round(sum(l["budget"] for l in lignes), 2),
        "total_engage": round(sum(l["engage"] for l in lignes), 2),
        "impute": round(float(subvention.total_impute_affectations or 0), 2),
        "reste_imputable": round(float(subvention.reste_imputable or 0), 2),
    }


def _journal_faits_marquants(subvention: Subvention, annee: int) -> list[ProjetJournalEntry]:
    projet_ids = [
        lien.projet_id for lien in (subvention.projets or []) if lien.projet_id
    ]
    if not projet_ids:
        return []
    return (
        ProjetJournalEntry.query
        .filter(
            ProjetJournalEntry.projet_id.in_(projet_ids),
            ProjetJournalEntry.entry_date >= date(annee, 1, 1),
            ProjetJournalEntry.entry_date <= date(annee, 12, 31),
        )
        .order_by(ProjetJournalEntry.entry_date.asc())
        .all()
    )


def _indicateurs_projets(subvention: Subvention, annee: int) -> list[dict]:
    from app.services.indicators import compute_project_indicators

    rows: list[dict] = []
    for lien in subvention.projets or []:
        projet = lien.projet
        if projet is None:
            continue
        for indicateur in compute_project_indicators(projet, selected_annee=annee):
            rows.append({"projet": projet.nom, **indicateur})
    return rows


def _paragraphes(subvention: Subvention, annee: int, stats: dict,
                 demo: dict | None, finances: dict) -> list[dict]:
    """Les paragraphes proposés : chiffres réels dans des formulations types.

    Chaque paragraphe indique sa source pour que l'équipe puisse vérifier.
    Rien n'est inventé : sans donnée, pas de paragraphe.
    """
    paras: list[dict] = []
    financeur = (subvention.financeur or "le financeur").strip()

    if stats["seances"]:
        paras.append({
            "titre": "Réalisation de l'action",
            "texte": (
                f"En {annee}, l'action « {subvention.nom} » soutenue par {financeur} "
                f"a représenté {_fmt_nombre(stats['seances'])} séances, soit "
                f"{_fmt_nombre(stats['heures'])} heures d'animation en face-à-face. "
                f"Elle a généré {_fmt_nombre(stats['participations'])} participations "
                f"pour {_fmt_nombre(stats['uniques'])} habitants différents."
            ),
            "source": "Émargement des ateliers financés (automatique).",
        })

    if demo:
        qpv = demo.get("qpv", {})
        genre = demo.get("genre", {})
        femmes = sum(n for g, n in genre.items() if (g or "").lower().startswith("f"))
        morceaux = []
        if demo.get("age_avg") is not None:
            morceaux.append(f"l'âge moyen des participants est de {demo['age_avg']} ans")
        if femmes:
            morceaux.append(f"{_fmt_nombre(femmes)} femmes ont été touchées")
        if qpv.get("qpv"):
            morceaux.append(
                f"{_fmt_nombre(qpv['qpv'])} participants résident en quartier "
                "prioritaire de la politique de la ville (QPV)"
            )
        if morceaux:
            texte = "Concernant les publics : " + " ; ".join(morceaux) + "."
            paras.append({
                "titre": "Publics touchés",
                "texte": texte,
                "source": "Fiches participants des présents (automatique).",
            })

    base_financiere = float(subvention.montant_attribue or subvention.montant_recu or 0)
    if base_financiere > 0 and stats["uniques"]:
        cout_participant = base_financiere / stats["uniques"]
        morceau_heure = ""
        if stats["heures"]:
            morceau_heure = f", et de {_fmt_euros(base_financiere / stats['heures'])} par heure d'animation"
        paras.append({
            "titre": "Coût unitaire",
            "texte": (
                f"Rapportée aux {_fmt_nombre(stats['uniques'])} habitants touchés, la "
                f"subvention de {_fmt_euros(base_financiere)} représente un coût de "
                f"{_fmt_euros(cout_participant)} par participant{morceau_heure}."
            ),
            "source": "Montant attribué ÷ activité réelle (automatique).",
        })

    if finances["total_budget"] > 0:
        taux = round(finances["total_engage"] / finances["total_budget"] * 100)
        paras.append({
            "titre": "Exécution budgétaire",
            "texte": (
                f"Sur un budget de {_fmt_euros(finances['total_budget'])}, "
                f"{_fmt_euros(finances['total_engage'])} ont été engagés à ce jour, "
                f"soit un taux d'exécution de {taux} %. Le détail par ligne figure "
                "dans le tableau financier joint."
            ),
            "source": "Lignes budgétaires et dépenses imputées (automatique).",
        })

    return paras


def _checklist(subvention: Subvention, stats: dict, finances: dict,
               narratif: BilanLourdNarratif | None) -> list[dict]:
    """Ce qui manque encore pour un bilan solide — avec où le corriger."""
    points = [
        {
            "label": "Ateliers financés reliés à la subvention",
            "ok": bool(stats["atelier_ids"]),
            "aide": "Sans ce lien, aucun chiffre d'activité ne peut être calculé (fiche subvention → ateliers financés).",
        },
        {
            "label": "Présences saisies sur l'exercice",
            "ok": stats["participations"] > 0,
            "aide": "Les chiffres viennent de l'émargement : vérifiez que les séances de l'année sont émargées.",
        },
        {
            "label": "Budget réel renseigné (lignes de charges)",
            "ok": finances["total_budget"] > 0,
            "aide": "Les lignes réelles servent de base au tableau financier.",
        },
        {
            "label": "Dépenses imputées sur la subvention",
            "ok": finances["impute"] > 0 or finances["total_engage"] > 0,
            "aide": "Les dépenses prouvent l'utilisation des fonds.",
        },
        {
            "label": "Narratif annuel du secteur rédigé",
            "ok": bool(narratif and (narratif.faits_marquants or narratif.difficultes or narratif.perspectives)),
            "aide": "Bilans → Bilans complets → champs libres : faits marquants, difficultés, perspectives.",
        },
    ]
    return points


def construire_trame(subvention: Subvention, annee: int | None = None) -> dict:
    """Assemble la trame complète du bilan financeur d'une subvention."""
    annee = int(annee or subvention.annee_exercice or date.today().year)

    stats = _stats_activite(subvention, annee)
    demo = _demographie(stats["atelier_ids"], annee)
    finances = _finances(subvention)
    narratif = BilanLourdNarratif.query.filter_by(
        annee=annee, secteur=subvention.secteur or ""
    ).first()

    return {
        "annee": annee,
        "subvention": subvention,
        "identification": {
            "nom": subvention.nom,
            "financeur": subvention.financeur,
            "reference": subvention.reference,
            "secteur": subvention.secteur,
            "statut": subvention.statut_label,
            "montant_demande": float(subvention.montant_demande or 0),
            "montant_attribue": float(subvention.montant_attribue or 0),
            "montant_recu": float(subvention.montant_recu or 0),
            "date_bilan_prevu": subvention.date_bilan_prevu,
        },
        "ateliers_finances": [
            {
                "nom": lien.atelier.nom,
                "poids_pct": float(lien.poids_pct or 100),
                "justification": lien.justification,
            }
            for lien in (subvention.ateliers_finances or [])
            if lien.atelier is not None and not lien.atelier.is_deleted
        ],
        "projets": [
            lien.projet.nom for lien in (subvention.projets or []) if lien.projet is not None
        ],
        "stats": stats,
        "demographie": demo,
        "finances": finances,
        "paragraphes": _paragraphes(subvention, annee, stats, demo, finances),
        "journal": _journal_faits_marquants(subvention, annee),
        "narratif_secteur": narratif,
        "indicateurs": _indicateurs_projets(subvention, annee),
        "checklist": _checklist(subvention, stats, finances, narratif),
    }


# ---------------------------------------------------------------------------
# Export DOCX éditable
# ---------------------------------------------------------------------------

def construire_docx(trame: dict):
    """Document Word éditable reprenant toute la trame (à finaliser à la main)."""
    from docx import Document

    doc = Document()
    ident = trame["identification"]
    annee = trame["annee"]

    doc.add_heading(f"Bilan {annee} — {ident['nom']}", level=1)
    doc.add_paragraph(
        f"Financeur : {ident['financeur'] or 'à compléter'} · "
        f"Référence : {ident['reference'] or 'à compléter'} · "
        f"Secteur : {ident['secteur'] or '—'}"
    )
    doc.add_paragraph(
        f"Montant demandé : {_fmt_euros(ident['montant_demande'])} · "
        f"attribué : {_fmt_euros(ident['montant_attribue'])} · "
        f"reçu : {_fmt_euros(ident['montant_recu'])}"
    )

    if trame["ateliers_finances"]:
        doc.add_heading("Action financée", level=2)
        for lien in trame["ateliers_finances"]:
            libelle = f"- {lien['nom']} ({lien['poids_pct']:.0f} % de la subvention)"
            if lien["justification"]:
                libelle += f" — {lien['justification']}"
            doc.add_paragraph(libelle)

    if trame["paragraphes"]:
        doc.add_heading("Éléments rédigés (à ajuster)", level=2)
        for para in trame["paragraphes"]:
            doc.add_heading(para["titre"], level=3)
            doc.add_paragraph(para["texte"])

    if trame["journal"]:
        doc.add_heading("Faits marquants (journal des projets liés)", level=2)
        for entree in trame["journal"]:
            prefixe = entree.entry_date.strftime("%d/%m/%Y")
            doc.add_paragraph(f"- {prefixe} — {entree.titre or ''} {entree.contenu}".strip())

    narratif = trame["narratif_secteur"]
    if narratif is not None:
        doc.add_heading("Matière issue du bilan annuel du secteur", level=2)
        for titre, contenu in (
            ("Faits marquants", narratif.faits_marquants),
            ("Difficultés rencontrées", narratif.difficultes),
            ("Perspectives", narratif.perspectives),
        ):
            if contenu:
                doc.add_heading(titre, level=3)
                doc.add_paragraph(contenu)

    finances = trame["finances"]
    if finances["lignes"]:
        doc.add_heading("Tableau financier", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        entetes = ("Compte", "Libellé", "Budget (€)", "Engagé (€)", "Reste (€)")
        for cellule, texte in zip(table.rows[0].cells, entetes):
            cellule.text = texte
        for ligne in finances["lignes"]:
            cells = table.add_row().cells
            cells[0].text = str(ligne["compte"] or "")
            cells[1].text = str(ligne["libelle"] or "")
            cells[2].text = f"{ligne['budget']:.2f}"
            cells[3].text = f"{ligne['engage']:.2f}"
            cells[4].text = f"{ligne['reste']:.2f}"
        totaux = table.add_row().cells
        totaux[1].text = "TOTAL"
        totaux[2].text = f"{finances['total_budget']:.2f}"
        totaux[3].text = f"{finances['total_engage']:.2f}"

    if trame["indicateurs"]:
        doc.add_heading("Indicateurs des projets liés", level=2)
        for row in trame["indicateurs"]:
            valeur = row.get("display_value") or row.get("value")
            cible = row.get("target")
            texte = f"- {row['projet']} · {row.get('label')} : {valeur if valeur is not None else 'non renseigné'}"
            if cible is not None:
                texte += f" (cible {row.get('target_symbol', '≥')} {cible})"
            doc.add_paragraph(texte)

    doc.add_paragraph()
    doc.add_paragraph(
        "Document généré automatiquement à partir des données de l'application "
        f"le {date.today().strftime('%d/%m/%Y')} — à relire et compléter avant envoi."
    )
    return doc
