"""Vie associative : instances de gouvernance, mandats, réunions, quorum.

Fonctions métier :
- comptage best-effort des adhérents à jour de cotisation (base d'un quorum d'AG) ;
- listes de travail : mandats arrivant à échéance, prochaines réunions ;
- alertes quotidiennes d'échéance de mandat (centre « À traiter ») ;
- génération d'une convocation et d'une feuille d'émargement au format Word.

Le quorum lui-même est porté par le modèle ``Reunion`` (présents + pouvoirs
vs. quorum requis). Ici on prépare la matière (base électeurs suggérée) et on
alerte ; on ne branche rien sur la comptabilité.
"""
from __future__ import annotations

from datetime import date

from flask import current_app

from app.extensions import db
from app.models import (
    MANDAT_ALERTE_JOURS,
    InstanceGouvernance,
    Mandat,
    Reunion,
    SuiviRappel,
)
from app.utils.dates import utcnow


class GouvernanceInvalide(ValueError):
    """Erreur métier à afficher telle quelle."""


# ---------------------------------------------------------------------------
# Adhérents à jour (base d'un quorum d'AG) — best-effort
# ---------------------------------------------------------------------------

def nombre_adherents_a_jour(annee_scolaire: int | None = None) -> int:
    """Compte, au mieux, les adhérents à jour de cotisation pour l'année :
    adhésions individuelles soldées (+1) et adhésions familiales soldées
    (+ tous les membres du foyer). Sert à *suggérer* la base d'électeurs d'une
    assemblée générale — jamais un chiffre imposé."""
    from app.models import Cotisation
    from app.services.cotisations import annee_scolaire_courante

    annee = annee_scolaire or annee_scolaire_courante()
    total = 0
    cotisations = Cotisation.query.filter(
        Cotisation.annee_scolaire == annee,
        Cotisation.type_cotisation.in_(("adhesion_individuelle", "adhesion_familiale")),
    ).all()
    for c in cotisations:
        if not c.solde:
            continue
        if c.type_cotisation == "adhesion_familiale" and c.foyer is not None:
            total += max(1, len(c.foyer.membres))
        else:
            total += 1
    return total


# ---------------------------------------------------------------------------
# Listes de travail
# ---------------------------------------------------------------------------

def instances_actives() -> list[InstanceGouvernance]:
    return (
        InstanceGouvernance.query.filter_by(actif=True)
        .order_by(InstanceGouvernance.type_instance.asc(), InstanceGouvernance.nom.asc())
        .all()
    )


def mandats_a_echeance(jours: int = MANDAT_ALERTE_JOURS) -> list[Mandat]:
    """Mandats actifs échus ou arrivant à échéance dans ``jours`` jours."""
    mandats = (
        Mandat.query.filter(Mandat.actif.is_(True), Mandat.date_fin.isnot(None))
        .order_by(Mandat.date_fin.asc())
        .all()
    )
    return [m for m in mandats if (m.jours_avant_echeance is not None and m.jours_avant_echeance <= jours)]


def prochaines_reunions(limit: int = 6) -> list[Reunion]:
    return (
        Reunion.query.filter(
            Reunion.statut == "planifiee",
            Reunion.date_reunion >= date.today(),
        )
        .order_by(Reunion.date_reunion.asc())
        .limit(limit)
        .all()
    )


# ---------------------------------------------------------------------------
# Réunions & émargement
# ---------------------------------------------------------------------------

def enregistrer_presence(reunion: Reunion, nom: str, etat: str,
                         represente_par: str | None = None) -> None:
    """Ajoute une ligne d'émargement. ``etat`` : present / pouvoir / excuse / absent."""
    from app.models import PresenceReunion

    nom = (nom or "").strip()
    if not nom:
        raise GouvernanceInvalide("Indiquez le nom de la personne.")
    db.session.add(PresenceReunion(
        reunion_id=reunion.id,
        nom=nom,
        present=(etat == "present"),
        pouvoir=(etat == "pouvoir"),
        excuse=(etat == "excuse"),
        represente_par=(represente_par or "").strip() or None if etat == "pouvoir" else None,
    ))
    db.session.commit()


# ---------------------------------------------------------------------------
# Alertes d'échéance de mandat (passe quotidienne)
# ---------------------------------------------------------------------------

def _lien(endpoint: str, **kwargs) -> str | None:
    from flask import url_for
    try:
        return url_for(endpoint, **kwargs)
    except Exception:
        return None


def alertes_mandats_du_jour() -> int:
    """Crée un rappel « À traiter » pour les mandats échus / proches non encore
    alertés (au plus une alerte / mandat / 30 jours). Retourne le nombre créé."""
    envoyes = 0
    mandats = Mandat.query.filter(Mandat.actif.is_(True), Mandat.date_fin.isnot(None)).all()
    for m in mandats:
        j = m.jours_avant_echeance
        if j is None or j > MANDAT_ALERTE_JOURS:
            continue
        if m.alerte_echeance_at is not None and (utcnow() - m.alerte_echeance_at).days < 30:
            continue
        instance_nom = m.instance.nom if m.instance else ""
        if m.est_echu:
            titre = f"Mandat échu — {m.nom} ({m.fonction_label})"
            corps = (
                f"Le mandat de {m.nom} — {m.fonction_label}, {instance_nom} — est arrivé "
                f"à échéance le {m.date_fin.strftime('%d/%m/%Y')}. Pensez à le renouveler."
            )
        else:
            titre = f"Mandat à renouveler — {m.nom} ({m.fonction_label})"
            corps = (
                f"Le mandat de {m.nom} — {m.fonction_label}, {instance_nom} — arrive à "
                f"échéance le {m.date_fin.strftime('%d/%m/%Y')} (dans {j} jour(s))."
            )
        db.session.add(SuiviRappel(
            titre=titre[:200], description=corps, categorie="gouvernance",
            priorite="warn", lien_url=_lien("gouvernance.instance_detail", instance_id=m.instance_id),
        ))
        m.alerte_echeance_at = utcnow()
        envoyes += 1
    if envoyes:
        db.session.commit()
    return envoyes


def alertes_gouvernance_quotidiennes_si_necessaire() -> None:
    """Passe quotidienne (au plus une fois/jour), calquée sur les rappels planning."""
    from app.models import TachePlanifiee

    try:
        tache = TachePlanifiee.query.filter_by(nom="gouvernance_alertes").first()
        if tache is not None and tache.derniere_execution is not None \
                and tache.derniere_execution.date() >= utcnow().date():
            return
        if tache is None:
            tache = TachePlanifiee(nom="gouvernance_alertes")
            db.session.add(tache)
        tache.derniere_execution = utcnow()
        db.session.commit()
        alertes_mandats_du_jour()
    except Exception:
        db.session.rollback()
        current_app.logger.exception("Échec des alertes de gouvernance quotidiennes")


# ---------------------------------------------------------------------------
# Documents Word : convocation + feuille d'émargement
# ---------------------------------------------------------------------------

def _org_nom() -> str:
    from app.services.instance_settings import resolve_identity
    defaut = current_app.config.get("APP_NAME") or "La structure"
    _, org, _, _ = resolve_identity(defaut, defaut)
    return org or defaut


def construire_convocation_docx(reunion: Reunion):
    from docx import Document

    doc = Document()
    doc.add_heading(f"Convocation — {reunion.type_label}", level=1)
    doc.add_paragraph(_org_nom())
    doc.add_paragraph(reunion.instance.nom if reunion.instance else "")
    doc.add_paragraph("")
    doc.add_paragraph(f"Objet : {reunion.titre}")
    ligne = f"Date : {reunion.date_reunion.strftime('%d/%m/%Y')}"
    if reunion.heure:
        ligne += f" à {reunion.heure}"
    doc.add_paragraph(ligne)
    if reunion.lieu:
        doc.add_paragraph(f"Lieu : {reunion.lieu}")

    doc.add_heading("Ordre du jour", level=2)
    lignes = [l.strip() for l in (reunion.ordre_du_jour or "").splitlines() if l.strip()]
    if lignes:
        for l in lignes:
            doc.add_paragraph(l, style="List Bullet")
    else:
        doc.add_paragraph("(à compléter)")

    if reunion.type_reunion in ("ag_ordinaire", "ag_extraordinaire"):
        doc.add_paragraph()
        doc.add_paragraph(
            "En cas d'empêchement, vous pouvez donner pouvoir à un·e autre "
            "membre pour vous représenter et voter en votre nom."
        )
    doc.add_paragraph()
    doc.add_paragraph(f"Fait le {date.today().strftime('%d/%m/%Y')}.")
    return doc


def construire_emargement_docx(reunion: Reunion):
    """Feuille d'émargement : une ligne par personne, colonne signature."""
    from docx import Document

    doc = Document()
    doc.add_heading(f"Feuille d'émargement — {reunion.type_label}", level=1)
    doc.add_paragraph(_org_nom())
    doc.add_paragraph(reunion.instance.nom if reunion.instance else "")
    doc.add_paragraph(f"{reunion.titre} — {reunion.date_reunion.strftime('%d/%m/%Y')}"
                      + (f" à {reunion.heure}" if reunion.heure else ""))

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, entete in enumerate(("Nom", "Présent·e / Pouvoir", "Représenté·e par", "Signature")):
        table.cell(0, i).text = entete

    presences = sorted(reunion.presences, key=lambda p: (p.nom or "").lower())
    for p in presences:
        cells = table.add_row().cells
        cells[0].text = p.nom or ""
        cells[1].text = p.statut_label
        cells[2].text = p.represente_par or ""
        cells[3].text = ""
    # Quelques lignes vierges pour les arrivants non pré-inscrits.
    for _ in range(max(0, 8 - len(presences))):
        cells = table.add_row().cells
        for c in cells:
            c.text = ""

    doc.add_paragraph()
    total = reunion.total_votes
    ligne = f"Présents : {reunion.nb_presents} · Pouvoirs : {reunion.nb_pouvoirs} · Total voix : {total}"
    if reunion.quorum_requis:
        ligne += f" · Quorum requis : {reunion.quorum_requis} → " + \
            ("ATTEINT" if reunion.quorum_atteint else "NON ATTEINT")
    doc.add_paragraph(ligne)
    return doc
