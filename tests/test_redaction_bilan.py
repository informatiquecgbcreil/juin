"""Aide à la rédaction du bilan financeur.

- la trame assemble les chiffres réels : activité des ateliers financés,
  coût unitaire, exécution budgétaire — rien d'inventé ;
- la checklist repère ce qui manque ;
- la page s'affiche avec les paragraphes et le bouton copier ;
- l'export DOCX contient la trame ;
- accès protégé.
"""
import uuid
from datetime import date

import pytest

ANNEE = 2001  # exercice bac à sable dédié à ce module


@pytest.fixture()
def contexte(app):
    """Une subvention avec atelier financé (1 séance de 2 h, 2 présents),
    une ligne de charge 1000 € dont 400 € dépensés, un narratif secteur."""
    suf = uuid.uuid4().hex[:6]
    with app.app_context():
        from app.extensions import db
        from app.models import (
            AtelierActivite,
            BilanLourdNarratif,
            Depense,
            LigneBudget,
            Participant,
            PresenceActivite,
            SessionActivite,
            Subvention,
            SubventionAtelier,
        )

        secteur = f"SecRed{suf}"
        sub = Subvention(
            nom=f"CLAS {suf}", secteur=secteur, annee_exercice=ANNEE,
            financeur="CAF Oise", reference=f"REF-{suf}",
            montant_demande=1200, montant_attribue=1000, montant_recu=600,
            date_bilan_prevu=date(ANNEE + 1, 1, 31),
        )
        atelier = AtelierActivite(nom=f"Accompagnement scolaire {suf}", secteur=secteur)
        p1 = Participant(nom=f"Red{suf}", prenom="Aya", genre="Femme",
                         date_naissance=date(ANNEE - 10, 3, 1))
        p2 = Participant(nom=f"Red{suf}", prenom="Bilal", genre="Homme",
                         date_naissance=date(ANNEE - 11, 6, 1))
        db.session.add_all([sub, atelier, p1, p2])
        db.session.flush()

        db.session.add(SubventionAtelier(subvention_id=sub.id, atelier_id=atelier.id,
                                         poids_pct=100, justification="Jeunes du QPV"))
        seance = SessionActivite(atelier_id=atelier.id, secteur=secteur,
                                 session_type="COLLECTIF", date_session=date(ANNEE, 3, 5),
                                 heure_debut="17:00", heure_fin="19:00", statut="realisee")
        ligne = LigneBudget(subvention_id=sub.id, nature="charge", compte="60",
                            libelle="Fournitures pédagogiques", montant_base=1000,
                            montant_reel=1000)
        db.session.add_all([seance, ligne])
        db.session.flush()
        db.session.add_all([
            PresenceActivite(session_id=seance.id, participant_id=p1.id),
            PresenceActivite(session_id=seance.id, participant_id=p2.id),
            Depense(ligne_budget_id=ligne.id, libelle="Cahiers", montant=400, statut="valide"),
        ])
        db.session.add(BilanLourdNarratif(
            annee=ANNEE, secteur=secteur,
            faits_marquants=f"Sortie musée très réussie {suf}",
        ))
        db.session.commit()
        return {"suf": suf, "subvention_id": sub.id, "secteur": secteur}


def test_trame_assemble_les_chiffres_reels(app, contexte):
    with app.app_context():
        from app.extensions import db
        from app.models import Subvention
        from app.services.redaction_bilan import construire_trame

        sub = db.session.get(Subvention, contexte["subvention_id"])
        trame = construire_trame(sub)

        assert trame["annee"] == ANNEE
        assert trame["stats"]["seances"] == 1
        assert trame["stats"]["participations"] == 2
        assert trame["stats"]["uniques"] == 2
        assert trame["stats"]["heures"] == 2.0

        assert trame["finances"]["total_budget"] == 1000.0
        assert trame["finances"]["total_engage"] == 400.0

        titres = {p["titre"] for p in trame["paragraphes"]}
        assert "Réalisation de l'action" in titres
        assert "Coût unitaire" in titres
        assert "Exécution budgétaire" in titres

        realisation = next(p for p in trame["paragraphes"] if p["titre"] == "Réalisation de l'action")
        assert "2 heures d'animation" in realisation["texte"]
        assert "2 habitants" in realisation["texte"]
        cout = next(p for p in trame["paragraphes"] if p["titre"] == "Coût unitaire")
        assert "500 €" in cout["texte"], "1000 € / 2 participants"
        budget = next(p for p in trame["paragraphes"] if p["titre"] == "Exécution budgétaire")
        assert "40 %" in budget["texte"]

        assert trame["narratif_secteur"] is not None
        assert contexte["suf"] in trame["narratif_secteur"].faits_marquants


def test_checklist_detecte_les_manques(app, contexte):
    with app.app_context():
        from app.extensions import db
        from app.models import Subvention
        from app.services.redaction_bilan import construire_trame

        # Subvention nue, sans atelier ni ligne : tout doit être signalé.
        nue = Subvention(nom=f"Nue {contexte['suf']}", secteur=contexte["secteur"],
                         annee_exercice=ANNEE)
        db.session.add(nue)
        db.session.commit()

        trame = construire_trame(nue)
        etats = {p["label"]: p["ok"] for p in trame["checklist"]}
        assert etats["Ateliers financés reliés à la subvention"] is False
        assert etats["Budget réel renseigné (lignes de charges)"] is False
        assert not trame["paragraphes"], "sans données, aucun paragraphe inventé"

        # La subvention complète du contexte coche tout.
        sub = db.session.get(Subvention, contexte["subvention_id"])
        etats = {p["label"]: p["ok"] for p in construire_trame(sub)["checklist"]}
        assert all(etats.values()), etats


def test_page_redaction(admin_client, contexte):
    r = admin_client.get(f"/bilans/redaction?subvention_id={contexte['subvention_id']}")
    assert r.status_code == 200
    page = r.get_data(as_text=True)
    assert "Rédiger le bilan financeur" in page
    assert f"CLAS {contexte['suf']}" in page
    # L'apostrophe est échappée en HTML (&#39;) : on teste sans elle.
    assert "Réalisation de l" in page
    assert "Copier" in page
    assert "Exporter la trame (Word)" in page


def test_export_docx(admin_client, contexte):
    r = admin_client.get(
        f"/bilans/redaction/export.docx?subvention_id={contexte['subvention_id']}"
    )
    assert r.status_code == 200
    assert "wordprocessingml" in r.content_type

    from io import BytesIO
    from docx import Document
    doc = Document(BytesIO(r.data))
    textes = "\n".join(p.text for p in doc.paragraphs)
    assert f"CLAS {contexte['suf']}" in textes
    assert "Réalisation de l'action" in textes
    assert "à relire et compléter avant envoi" in textes
    assert len(doc.tables) == 1, "le tableau financier est présent"


def test_refus_anonymes(client):
    assert client.get("/bilans/redaction").status_code in (302, 401, 403)
    assert client.get("/bilans/redaction/export.docx").status_code in (302, 401, 403)
